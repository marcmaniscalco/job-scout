import io
from unittest.mock import patch

from docx import Document

from job_scout.clients.s3_client import S3ResumeStore
from tests.conftest import BUCKET_NAME, REGION


def _build_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Backend Engineer")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "AWS"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_get_resume_text_fetches_and_decodes_object(resume_bucket):
    resume_bucket.put_object(
        Bucket=BUCKET_NAME, Key="resume.txt", Body=b"My resume text"
    )
    store = S3ResumeStore(BUCKET_NAME, "resume.txt", REGION)

    assert store.get_resume_text() == "My resume text"


def test_get_resume_text_caches_after_first_call(resume_bucket):
    resume_bucket.put_object(
        Bucket=BUCKET_NAME, Key="resume.txt", Body=b"My resume text"
    )
    store = S3ResumeStore(BUCKET_NAME, "resume.txt", REGION)

    with patch.object(
        store._client, "get_object", wraps=store._client.get_object
    ) as spy:
        store.get_resume_text()
        store.get_resume_text()

    assert spy.call_count == 1


def test_get_resume_text_extracts_docx_paragraphs_and_tables(
    resume_bucket,
):
    resume_bucket.put_object(
        Bucket=BUCKET_NAME, Key="resume.docx", Body=_build_docx_bytes()
    )
    store = S3ResumeStore(BUCKET_NAME, "resume.docx", REGION)

    text = store.get_resume_text()

    assert "Jane Doe" in text
    assert "Senior Backend Engineer" in text
    assert "Python" in text
    assert "AWS" in text
