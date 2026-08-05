"""Resume fetch from S3, cached for the lifetime of the store instance.

The Lambda handler holds a single instance per warm container (see
handler.get_resume_store), so the resume is fetched from S3 at most
once per container instead of once per invocation.
"""

import io

import boto3
from docx import Document


class S3ResumeStore:
    def __init__(
        self, bucket: str, object_key: str, region_name: str | None = None
    ):
        self._bucket = bucket
        self._object_key = object_key
        self._client = boto3.client("s3", region_name=region_name)
        self._cached_text: str | None = None

    def get_resume_text(self) -> str:
        if self._cached_text is None:
            response = self._client.get_object(
                Bucket=self._bucket, Key=self._object_key
            )
            content = response["Body"].read()
            if self._object_key.lower().endswith(".docx"):
                self._cached_text = self._extract_docx_text(content)
            else:
                self._cached_text = content.decode("utf-8")
        return self._cached_text

    @staticmethod
    def _extract_docx_text(content: bytes) -> str:
        document = Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs]
        # Resumes often use tables for layout (e.g. skills lists), so
        # pull cell text too rather than silently dropping it.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(p for p in paragraphs if p.strip())
